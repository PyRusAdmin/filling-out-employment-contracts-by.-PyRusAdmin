from loguru import logger

from src.database import read_from_db

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

datas = [
    21982, 347, 621, 6025, 8274, 20461, 21849, 22186, 22465, 22769, 23412, 467, 3379, 3549, 4451, 13700, 17821,
    20499, 21152, 22430, 1069, 1089, 22575, 4464, 5851, 6302, 3383, 6139, 9352, 22118, 2355, 4363, 5201,
    16744, 16831, 16937, 21812, 22273, 23419, 5867, 17646, 22524, 23270, 23306, 23456, 438, 10397,
    16861, 11441, 22271, 22272, 22958, 4232, 6106, 22256, 23341, 23399, 20236, 5538, 6963, 21703, 21844, 10893,
    23130, 20644, 22784, 22822, 22922, 23478, 10973, 11980, 15876, 22726, 13123, 15437, 21964, 22862, 23470, 7092,
    7742, 13754, 15750, 22652, 23144, 12000, 21227, 22928, 22819
]

async def address_parsing():
    """Парсинг адреса"""

    logger.info("Парсинг адреса")

    rows = await read_from_db()

    if not rows:
        logger.warning("Нет данных из БД")
        return []

    data = []
    matches = []

    for row in rows:
        raw_tab_num = row.a4_табельный_номер

        # Логируем "как есть"
        logger.debug(f"Raw табельный номер: {raw_tab_num}, тип: {type(raw_tab_num)}")

        # Преобразуем в int с обработкой ошибок
        try:
            tab_num = int(str(raw_tab_num).strip())
        except (ValueError, AttributeError, TypeError):
            logger.warning(f"Пропущено: не удалось обработать табельный номер '{raw_tab_num}'")
            continue

        # сохраняем данные как словарь
        entry = {
            "табельный": tab_num,
            "a5": row.a5,
            "a13": row.a13
        }
        data.append(entry)

        # проверяем совпадение
        if tab_num in datas:
            matches.append(entry)
            logger.debug(f"Найдено совпадение: {entry}")

    logger.info(f"Всего записей: {len(data)}")
    logger.info(f"Совпадений найдено: {len(matches)}")

    # Сохраняем в Word
    if matches:
        save_matches_to_docx(matches)
    else:
        logger.warning("Совпадений не найдено, файл не создан.")

    return matches


def save_matches_to_docx(matches, filename=None):
    """
    Сохраняет список совпадений в .docx файл.

    :param matches: список словарей с ключами "табельный", "a5", "a13"
    :param filename: имя файла (опционально)
    :return: путь к файлу
    """
    if not matches:
        print("Нет данных для сохранения.")
        return None

    # Создаём документ
    doc = Document()

    # Заголовок
    title = doc.add_heading('Совпадения табельных номеров', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Подзаголовок с датой
    current_date = datetime.now().strftime("%d.%m.%Y %H:%M")
    doc.add_paragraph(f"Дата формирования: {current_date}")
    doc.add_paragraph()

    # Создаём таблицу: 1 строка — заголовки, остальные — данные
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Заголовки
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Табельный номер'
    hdr_cells[1].text = 'A5'
    hdr_cells[2].text = 'A13'

    # Заполняем данными
    for match in matches:
        row_cells = table.add_row().cells
        row_cells[0].text = str(match["табельный"])
        row_cells[1].text = str(match["a5"]) if match["a5"] else ""
        row_cells[2].text = str(match["a13"]) if match["a13"] else ""

    # Настраиваем шрифт (опционально)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

    # Сохраняем
    if not filename:
        filename = f"Совпадения_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    # Можно указать папку, например, "output"
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True)
    filepath = os.path.join(output_dir, filename)

    doc.save(filepath)
    print(f"Данные сохранены в файл: {filepath}")
    return filepath